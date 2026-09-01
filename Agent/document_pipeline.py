"""
Phase 1 MVP — File -> Extract -> LLM pipeline.

Upload File -> Detect File Type -> Extract Content -> Convert to common text
-> User Question -> LLM -> Answer + source

Supported: PDF, DOCX, TXT, CSV, XLSX.

This module provides:
  - Extraction: turn an uploaded binary file into plain text.
  - Storage: in-memory document store (keyed by doc_id).
  - Generation: create PDF / DOCX / TXT files from content and save them to
    disk, inside the currently active agent working directory (see
    _generated_docs_dir / path_utils.current_workdir).
"""
import io
import os
import uuid
from datetime import datetime
from pathlib import Path

from path_utils import current_workdir

MAX_CHARS_FOR_CONTEXT = 12000  # MVP: whole-document context, no chunking/retrieval yet
MAX_IMAGE_BYTES = 8 * 1024 * 1024  # 8MB cap on a single image attachment

TEXT_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".xlsx"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | IMAGE_EXTENSIONS


def _generated_docs_dir() -> Path:
    """Where generate_pdf/generate_docx/generate_txt save their output --
    inside whichever working directory is active for the current run (the
    user's selected cowork folder, or DEFAULT_AGENT_WORKDIR if none was
    selected), never a fixed folder next to this module. Same pattern as
    path_utils.current_artifacts_dir()."""
    d = current_workdir() / "generated_documents"
    d.mkdir(parents=True, exist_ok=True)
    return d

_IMAGE_MIME_TYPES = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
}


class UnsupportedFileType(Exception):
    pass


def detect_file_type(filename: str) -> str:
    ext = os.path.splitext(filename or "")[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileType(
            f"Unsupported file type: '{ext or 'unknown'}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    return ext


def is_image_extension(ext: str) -> bool:
    return ext.lower() in IMAGE_EXTENSIONS


def image_mime_type(ext: str) -> str:
    return _IMAGE_MIME_TYPES.get(ext.lower(), "application/octet-stream")


# ---------- extractors (one per type, common text out) ----------

def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"[Page {i}]\n{text}")
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def _extract_csv(data: bytes) -> str:
    import pandas as pd
    df = pd.read_csv(io.BytesIO(data))
    return df.to_csv(index=False)


def _extract_xlsx(data: bytes) -> str:
    import pandas as pd
    sheets = pd.read_excel(io.BytesIO(data), sheet_name=None)
    parts = [f"[Sheet: {name}]\n{df.to_csv(index=False)}" for name, df in sheets.items()]
    return "\n\n".join(parts)


_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".txt": _extract_txt,
    ".csv": _extract_csv,
    ".xlsx": _extract_xlsx,
}


def extract_text(filename: str, data: bytes) -> str:
    """Detect file type, extract, and convert to a common plain-text format."""
    ext = detect_file_type(filename)
    if ext not in _EXTRACTORS:
        # ext is a valid SUPPORTED_EXTENSIONS member (e.g. an image type) but
        # has no text extractor -- fail with a clear message instead of a
        # raw KeyError if a caller ever routes an image here directly
        # (process_upload() itself avoids this by checking is_image_extension first).
        raise UnsupportedFileType(
            f"'{ext}' has no text extractor -- route image files through add_image() instead."
        )
    text = _EXTRACTORS[ext](data).strip()
    if not text:
        raise ValueError("No extractable text found in this file.")
    return text


# ---------- storage (in-memory, per-process — Phase 1 only) ----------

class DocumentStore:
    def __init__(self):
        self._docs: dict[str, dict] = {}

    def add(self, filename: str, text: str) -> str:
        doc_id = str(uuid.uuid4())
        self._docs[doc_id] = {
            "doc_id": doc_id,
            "filename": filename,
            "is_image": False,
            "text": text,
            "char_count": len(text),
            "uploaded_at": datetime.now().isoformat(),
        }
        return doc_id

    def add_image(self, filename: str, data: bytes) -> str:
        """Store an uploaded image as base64 for later multimodal attachment
        (vision), instead of forcing it through a text extractor. Kept
        separate from add() because images have no "text" -- callers that
        need one (read_document, prompt folding) get a short placeholder."""
        import base64

        ext = detect_file_type(filename)
        if not is_image_extension(ext):
            raise UnsupportedFileType(f"add_image called with a non-image file: {filename!r}")
        if len(data) > MAX_IMAGE_BYTES:
            raise ValueError(
                f"Image too large ({len(data)} bytes) -- limit is {MAX_IMAGE_BYTES} bytes."
            )

        doc_id = str(uuid.uuid4())
        self._docs[doc_id] = {
            "doc_id": doc_id,
            "filename": filename,
            "is_image": True,
            "mime_type": image_mime_type(ext),
            "data_b64": base64.b64encode(data).decode("ascii"),
            "text": f"[Image file: {filename}]",
            "char_count": 0,
            "uploaded_at": datetime.now().isoformat(),
        }
        return doc_id

    def get(self, doc_id: str) -> dict | None:
        return self._docs.get(doc_id)

    def list(self) -> list[dict]:
        return [
            {k: v for k, v in d.items() if k not in ("text", "data_b64")}
            for d in sorted(self._docs.values(), key=lambda d: d["uploaded_at"], reverse=True)
        ]


document_store = DocumentStore()


def process_upload(filename: str, data: bytes) -> str:
    """Unified upload entrypoint: routes to text extraction or image storage
    based on extension, and returns the resulting doc_id either way -- the
    one function an upload endpoint needs to call regardless of file kind."""
    ext = detect_file_type(filename)
    if is_image_extension(ext):
        return document_store.add_image(filename, data)
    text = extract_text(filename, data)
    return document_store.add(filename, text)


# ---------- prompt helper (for folding doc text into user messages) ----------

def build_prompt_with_document(user_text: str, doc_id: str) -> str:
    """Fold a stored document's extracted text into the user's message content so it
    flows through the normal orchestrator pipeline (planner reads state["messages"][-1].content)
    instead of being answered by a separate, isolated LLM call."""
    doc = document_store.get(doc_id)
    if doc is None:
        raise KeyError(f"No document found with id: {doc_id}")

    context = doc["text"][:MAX_CHARS_FOR_CONTEXT]
    truncated = len(doc["text"]) > MAX_CHARS_FOR_CONTEXT
    note = " (truncated)" if truncated else ""

    return (
        f"{user_text}\n\n"
        f"--- Attached document: {doc['filename']}{note} ---\n"
        f"{context}\n"
        f"--- end document ---"
    )


def build_multimodal_message(user_text: str, doc_ids: list[str] | None) -> str | list[dict]:
    """Fold a mix of attached documents/images into one message payload.

    Text-extractable docs (pdf/docx/txt/csv/xlsx) are folded into the text
    the same way build_prompt_with_document does. Images become separate
    OpenAI-style `image_url` content blocks so a vision-capable model can
    actually see them -- something a single text string can never carry.

    Returns a plain str when there are no images (so callers that don't
    care about attachments can keep treating the result as ordinary text),
    or a list of content blocks (str block + image blocks) when at least
    one image is attached -- the shape LangChain's HumanMessage(content=...)
    expects for multimodal input.
    """
    if not doc_ids:
        return user_text

    text_parts = [user_text]
    image_blocks: list[dict] = []
    for doc_id in doc_ids:
        doc = document_store.get(doc_id)
        if doc is None:
            text_parts.append(f"\n\n[Attachment error: no document found with id {doc_id}]")
            continue
        if doc.get("is_image"):
            image_blocks.append(
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{doc['mime_type']};base64,{doc['data_b64']}"},
                }
            )
            text_parts.append(f"\n\n[Attached image: {doc['filename']}]")
        else:
            context = doc["text"][:MAX_CHARS_FOR_CONTEXT]
            truncated = len(doc["text"]) > MAX_CHARS_FOR_CONTEXT
            note = " (truncated)" if truncated else ""
            text_parts.append(
                f"\n\n--- Attached document: {doc['filename']}{note} ---\n{context}\n--- end document ---"
            )

    if not image_blocks:
        return "".join(text_parts)

    return [{"type": "text", "text": "".join(text_parts)}, *image_blocks]


# ---------- document generators ----------

def _unique_filename(basename: str, ext: str) -> str:
    """Generate a collision-free filename like 'report_a1b2c3d4.pdf'."""
    stem = os.path.splitext(basename)[0] if basename else "document"
    # Sanitise: keep only alphanumeric, dash, underscore
    stem = "".join(c if c.isalnum() or c in "-_" else "_" for c in stem).strip("_") or "document"
    short_id = uuid.uuid4().hex[:8]
    return f"{stem}_{short_id}{ext}"


def generate_pdf(content: str, filename: str = "document") -> dict:
    """Generate a PDF file from text/markdown content using fpdf2.

    Returns dict with filepath, filename, and download_url on success,
    or an error dict on failure.
    """
    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)

        def sanitize_latin1(text: str) -> str:
            replacements = {
                '\u2018': "'", '\u2019': "'",
                '\u201c': '"', '\u201d': '"',
                '\u2013': '-', '\u2014': '--',
                '\u2026': '...',
                '\u2022': '-',
                '\u00a0': ' ',
                '\u2713': 'v',
                '\u2714': 'v',
            }
            for k, v in replacements.items():
                text = text.replace(k, v)
            # FPDF default Helvetica font uses latin-1
            return text.encode('latin-1', 'replace').decode('latin-1')

        for line in content.split("\n"):
            stripped = sanitize_latin1(line.strip())
            line_sanitized = sanitize_latin1(line)

            # Heading detection (markdown '#')
            if stripped.startswith("#"):
                level = len(stripped) - len(stripped.lstrip("#"))
                text = stripped.lstrip("#").strip()
                font_size = max(24 - (level - 1) * 4, 12)
                pdf.set_font("Helvetica", "B", font_size)
                pdf.cell(0, font_size * 0.6, text, new_x="LMARGIN", new_y="NEXT")
                pdf.ln(2)
                pdf.set_font("Helvetica", size=12)

            # Bullet list
            elif stripped.startswith(("-", "*")) and len(stripped) > 1 and stripped[1] == " ":
                text = stripped[2:].strip()
                pdf.cell(10)  # indent
                pdf.cell(0, 7, f"-  {text}", new_x="LMARGIN", new_y="NEXT")

            # Empty line
            elif not stripped:
                pdf.ln(5)

            # Normal text
            else:
                pdf.multi_cell(0, 7, line_sanitized)

        out_name = _unique_filename(filename, ".pdf")
        out_path = str(_generated_docs_dir() / out_name)
        pdf.output(out_path)

        return {
            "status": "success",
            "filename": out_name,
            "filepath": out_path,
            # download_url intentionally omitted: serving it is now
            # /generated-docs/{thread_id}/{filename} (see main.py), and this
            # module operates below the thread/HTTP layer -- it has no
            # thread_id to build that URL with. The caller (main.py's
            # /generated-docs listing, or whoever has thread_id) builds it.
            "download_url": None,
        }
    except Exception as e:
        return {"status": "error", "error": f"Failed to generate PDF: {e}"}


def generate_docx(content: str, filename: str = "document") -> dict:
    """Generate a DOCX file from text/markdown content using python-docx.

    Returns dict with filepath, filename, and download_url on success.
    """
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()

        for line in content.split("\n"):
            stripped = line.strip()

            # Heading detection
            if stripped.startswith("#"):
                level = len(stripped) - len(stripped.lstrip("#"))
                text = stripped.lstrip("#").strip()
                heading_level = min(level, 4)  # Word supports heading levels 1-9
                doc.add_heading(text, level=heading_level)

            # Bullet list
            elif stripped.startswith(("-", "*")) and len(stripped) > 1 and stripped[1] == " ":
                text = stripped[2:].strip()
                doc.add_paragraph(text, style="List Bullet")

            # Empty line — skip (Word handles paragraph spacing)
            elif not stripped:
                continue

            # Normal text
            else:
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.size = Pt(11)

        out_name = _unique_filename(filename, ".docx")
        out_path = str(_generated_docs_dir() / out_name)
        doc.save(out_path)

        return {
            "status": "success",
            "filename": out_name,
            "filepath": out_path,
            # See the matching comment in generate_pdf -- this module has no
            # thread_id, so it can't build the /generated-docs/{thread_id}/...
            # URL itself.
            "download_url": None,
        }
    except Exception as e:
        return {"status": "error", "error": f"Failed to generate DOCX: {e}"}


def generate_txt(content: str, filename: str = "document") -> dict:
    """Save content as a plain-text file.

    Returns dict with filepath, filename, and download_url on success.
    """
    try:
        out_name = _unique_filename(filename, ".txt")
        out_path = str(_generated_docs_dir() / out_name)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(content)

        return {
            "status": "success",
            "filename": out_name,
            "filepath": out_path,
            # See the matching comment in generate_pdf -- this module has no
            # thread_id, so it can't build the /generated-docs/{thread_id}/...
            # URL itself.
            "download_url": None,
        }
    except Exception as e:
        return {"status": "error", "error": f"Failed to generate TXT: {e}"}


GENERATORS = {
    "pdf": generate_pdf,
    "docx": generate_docx,
    "txt": generate_txt,
}


def generate_document(content: str, filename: str = "document", fmt: str = "pdf") -> dict:
    """Unified entry point for document generation.

    Parameters
    ----------
    content : str
        The text/markdown content to write into the document.
    filename : str
        Base name for the output file (extension is added automatically).
    fmt : str
        Output format: 'pdf', 'docx', or 'txt'.

    Returns
    -------
    dict  with 'status', 'filename', 'filepath', 'download_url' on success.
    """
    fmt = (fmt or "pdf").lower().strip().lstrip(".")
    generator = GENERATORS.get(fmt)
    if generator is None:
        return {
            "status": "error",
            "error": f"Unsupported format: '{fmt}'. Supported: {', '.join(sorted(GENERATORS))}",
        }
    return generator(content, filename)